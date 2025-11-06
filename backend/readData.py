import os
from docx import Document
from pydantic import BaseModel, Field
from typing import List, Optional


# -------------------------------
# 1️⃣ Define the structured schema
# -------------------------------

class FieldItem(BaseModel):
    field_text: str = Field(..., description="Line or field text possibly containing blanks or placeholders.")
    context: Optional[str] = Field(None, description="Nearby context to help LLM fill the value.")


class TableCell(BaseModel):
    row_index: int
    col_index: int
    cell_text: str
    context: Optional[str] = None


class PageStructure(BaseModel):
    page_number: int
    lines: List[FieldItem] = []
    tables: List[List[TableCell]] = []


class DocumentStructure(BaseModel):
    file_name: str
    pages: List[PageStructure]


# -------------------------------
# 2️⃣ Extract structure from DOCX
# -------------------------------

def read_docx_structure(docx_path: str) -> DocumentStructure:
    """Reads a DOCX file and extracts structured page-wise text + tables."""
    print(f"📄 Reading: {docx_path}")

    doc = Document(docx_path)
    file_name = os.path.basename(docx_path)
    pages: List[PageStructure] = []

    # Note: python-docx doesn’t detect real 'pages' (Word doesn’t store page breaks as layout info)
    # so we'll simulate pages by splitting on explicit page breaks or count-based heuristics
    page_number = 1
    current_page = PageStructure(page_number=page_number, lines=[], tables=[])

    for element in doc.element.body:
        if element.tag.endswith("p"):  # Paragraph
            text = element.text.strip()
            if text:
                current_page.lines.append(FieldItem(field_text=text))
        elif element.tag.endswith("tbl"):  # Table
            table_obj = []
            table = element
            # Handle merged or split tables as continuous blocks
            for r_idx, row in enumerate(table.findall(".//w:tr", table.nsmap)):
                for c_idx, cell in enumerate(row.findall(".//w:tc", table.nsmap)):
                    cell_text = "".join(t.text or "" for t in cell.findall(".//w:t", table.nsmap)).strip()
                    table_obj.append(TableCell(row_index=r_idx, col_index=c_idx, cell_text=cell_text))
            current_page.tables.append(table_obj)

        # Detect manual page breaks
        if element.tag.endswith("p"):
            for br in element.findall(".//w:br", element.nsmap):
                if br.attrib.get(f"{{{element.nsmap['w']}}}type") == "page":
                    pages.append(current_page)
                    page_number += 1
                    current_page = PageStructure(page_number=page_number, lines=[], tables=[])

    # Append last page
    if current_page.lines or current_page.tables:
        pages.append(current_page)

    return DocumentStructure(file_name=file_name, pages=pages)


# -------------------------------
# 3️⃣ Example usage
# -------------------------------
import json
if __name__ == "__main__":
    file_path = "filled_Lease Abstract Template.docx"  # change to your file
    structured_doc = read_docx_structure(file_path)
    
    with open("structure.json", "w", encoding="utf-8") as f:
        json.dump(structured_doc.model_dump(), f, ensure_ascii=False, indent=2)

    print("✅ Extracted structure:")
    for page in structured_doc.pages:
        print(f"\n📘 Page {page.page_number}")
        for line in page.lines:
            print(f"  - {line.field_text[:80]}")
        for t_idx, table in enumerate(page.tables):
            print(f"  🧾 Table {t_idx + 1} has {len(table)} cells")

    # Later: pass this structured_doc to your llm_structured() for filling
