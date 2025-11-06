import os
import json
from typing import List, Optional, Union
from pydantic import BaseModel, Field
from docx import Document
from openpyxl import load_workbook


# ===========================================================
# 1️⃣  Define unified structure model for both DOCX and XLSX
# ===========================================================

class FieldItem(BaseModel):
    field_text: str = Field(..., description="Raw text of the field or label, e.g., 'Tenant Name:'")
    value: Optional[str] = Field(None, description="Value placeholder (usually blank initially)")
    context: Optional[str] = Field(None, description="Nearby context or section heading if available")
    position: Optional[List[int]] = Field(None, description="[x, y] coordinates or line index in doc")


class TableCell(BaseModel):
    row_index: int
    col_index: int
    cell_text: str
    context: Optional[str] = None


class TableStructure(BaseModel):
    table_name: Optional[str] = None
    headers: List[str] = Field(default_factory=list)
    rows: List[List[TableCell]] = Field(default_factory=list)
    context: Optional[str] = None


class PageStructure(BaseModel):
    page_number: int
    fields: List[FieldItem] = Field(default_factory=list)
    tables: List[TableStructure] = Field(default_factory=list)


class DocumentStructure(BaseModel):
    file_name: str
    file_type: str
    pages: List[PageStructure]


# ===========================================================
# 2️⃣  DOCX Extraction Logic
# ===========================================================

def extract_from_docx(file_path: str) -> DocumentStructure:
    doc = Document(file_path)
    file_name = os.path.basename(file_path)
    print(f"📘 Extracting from DOCX → {file_name}")

    pages: List[PageStructure] = []
    page_number = 1
    current_page = PageStructure(page_number=page_number, fields=[], tables=[])

    for block in doc.element.body:
        tag = block.tag.split('}')[-1]

        # --- Paragraph / Text block ---
        if tag == "p":
            text = block.text.strip()
            if text:
                # Identify field-like lines (heuristic: contains ':' or underscores)
                if ":" in text or "___" in text:
                    current_page.fields.append(FieldItem(field_text=text))
                else:
                    # Use as context if previous field or table exists
                    if current_page.fields:
                        current_page.fields[-1].context = (
                            (current_page.fields[-1].context or "") + " " + text
                        ).strip()

        # --- Table ---
        # elif tag == "tbl":
        #     table = []
        #     for r_idx, row in enumerate(block.findall(".//w:tr", block.nsmap)):
        #         row_cells = []
        #         for c_idx, cell in enumerate(row.findall(".//w:tc", block.nsmap)):
        #             cell_text = "".join(
        #                 t.text or "" for t in cell.findall(".//w:t", block.nsmap)
        #             ).strip()
        #             row_cells.append(TableCell(row_index=r_idx, col_index=c_idx, cell_text=cell_text))
        #         table.append(row_cells)

        #     # Infer headers
        #     headers = [c.cell_text for c in table[0]] if table else []
        #     current_page.tables.append(
        #         TableStructure(headers=headers, rows=table)
        #     )
        elif tag == "tbl":
            table = []
            for r_idx, row in enumerate(block.findall(".//w:tr", block.nsmap)):
                row_cells = []
                for c_idx, cell in enumerate(row.findall(".//w:tc", block.nsmap)):
                    cell_text = "".join(
                        t.text or "" for t in cell.findall(".//w:t", block.nsmap)
                    ).strip()
                    row_cells.append(
                        TableCell(row_index=r_idx, col_index=c_idx, cell_text=cell_text)
                    )
                table.append(row_cells)

            # Default values
            table_label = None
            headers = []
            data_rows = table

            # Detect if first row is a label row (only one non-empty cell)
            if table:
                first_row_texts = [c.cell_text for c in table[0]]
                non_empty_count = sum(1 for text in first_row_texts if text.strip())

                if non_empty_count == 1 and len(table) > 1:
                    # First row is label row
                    table_label = next(text for text in first_row_texts if text.strip())
                    headers = [c.cell_text for c in table[1]]
                    data_rows = table[2:]
                else:
                    # Normal table
                    headers = [c.cell_text for c in table[0]]
                    data_rows = table[1:]

            current_page.tables.append(
                TableStructure(
                    table_name=table_label,
                    headers=headers,
                    rows=data_rows
                )
            )


        # Detect manual page breaks
        for br in block.findall(".//w:br", block.nsmap):
            if br.attrib.get(f"{{{block.nsmap['w']}}}type") == "page":
                pages.append(current_page)
                page_number += 1
                current_page = PageStructure(page_number=page_number, fields=[], tables=[])

    # Append last page
    if current_page.fields or current_page.tables:
        pages.append(current_page)

    return DocumentStructure(file_name=file_name, file_type="docx", pages=pages)


# ===========================================================
# 3️⃣  Excel Extraction Logic
# ===========================================================

def extract_from_excel(file_path: str) -> DocumentStructure:
    wb = load_workbook(file_path, data_only=True)
    file_name = os.path.basename(file_path)
    print(f"📗 Extracting from Excel → {file_name}")

    pages: List[PageStructure] = []
    page_number = 1

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        page = PageStructure(page_number=page_number, fields=[], tables=[])
        table = TableStructure(table_name=sheet, headers=[], rows=[])

        # Extract header row
        first_row = [cell.value for cell in ws[1]]
        table.headers = [str(h).strip() if h else "" for h in first_row]

        for r_idx, row in enumerate(ws.iter_rows(values_only=True)):
            row_cells = []
            for c_idx, val in enumerate(row):
                text = str(val).strip() if val else ""
                row_cells.append(TableCell(row_index=r_idx, col_index=c_idx, cell_text=text))
            table.rows.append(row_cells)

        page.tables.append(table)
        pages.append(page)
        page_number += 1

    return DocumentStructure(file_name=file_name, file_type="xlsx", pages=pages)


# ===========================================================
# 4️⃣  Router — Detect and Parse File Type
# ===========================================================

def extract_structure(file_path: str, output_path: str = "structure.json"):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        structure = extract_from_docx(file_path)
    elif ext in [".xls", ".xlsx"]:
        structure = extract_from_excel(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(structure.model_dump(), f, ensure_ascii=False, indent=2)

    print(f"✅ Structure extracted → {output_path}")
    print(f"   Pages: {len(structure.pages)}")
    return structure


# ===========================================================
# 5️⃣  Example Usage
# ===========================================================

if __name__ == "__main__":
    # file_path = "Lease Abstract Template.docx"  # or Excel file
    # file_path = "userPdfData/refinalfollowupactionrequiredforyourprototypedeve (1)/BOV Template.xlsx"
    file_path = "userPdfData/refinalfollowupactionrequiredforyourprototypedeve (1)/Lease Abstract Template.docx"
    try:
        structure = extract_structure(file_path)
    except Exception as e:
        print(f"⚠️ Error: {e}")
