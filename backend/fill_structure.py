import json
import random
import shutil
import traceback
import re
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt
from pydantic import BaseModel, Field
from noPklRetrieval import create_retriever, rag
from ragAnything import query_rag

# ==========================================
# 🔹 Pydantic Models
# ==========================================
class ExtractedField(BaseModel):
    value: str = Field(..., description="Extracted value")
    confidence: float = Field(..., description="Confidence 0-1")
    source: str = Field(..., description="Chunk or file source name")
    reason: str = Field(..., description="Why this value was chosen by LLM")


# class ExtractedTable(BaseModel):
#     rows: List[Dict[str, str]] = Field(..., description="List of rows in the table")
#     confidence: float = Field(..., description="Overall confidence between 0 and 1")
#     source: str = Field(..., description="Document chunk or source name")

# class ExtractedTable(BaseModel):
#     name: str = Field(..., description="Table name or title")
#     headers: List[str] = Field(..., description="Column headers for the table")
#     rows: List[Dict[str, str]] = Field(..., description="List of table rows with column-value mapping")
#     confidence: float = Field(..., description="Confidence level of the table extraction, 0-1")
#     source: str = Field(..., description="Source document or chunk name")
#     reason: Optional[str] = Field(None, description="Reasoning or explanation for extracted values")

#     class Config:
#         extra = "forbid"  # ensures additionalProperties=false

class ExtractedTable(BaseModel):
    """Represents a structured table extracted from a document."""

    name: str = Field(..., description="Table name or title")
    headers: List[str] = Field(..., description="Column headers for the table")
    rows: List[List[str]] = Field(
        ..., description="List of table rows, each a list of cell values in header order"
    )
    confidence: float = Field(..., description="Confidence level of the table extraction, between 0 and 1")
    source: str = Field(..., description="Source document or chunk name")
    reason: Optional[str] = Field(None, description="Reasoning or explanation for extracted values")

    model_config = {
        "extra": "forbid",
        "json_schema_extra": {
            # OpenAI’s structured output requires this
            "type": "object",
            "required": ["name", "headers", "rows", "confidence", "source"],
        },
    }


# ==========================================
# 🔹 Text Normalization Helpers
# ==========================================
def normalize(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip()).lower()

def tokens(s: str) -> List[str]:
    return [t for t in re.split(r"[^a-z0-9]+", normalize(s)) if t]

def jaccard(a: List[str], b: List[str]) -> float:
    A, B = set(a), set(b)
    if not (A or B):
        return 0.0
    return len(A & B) / len(A | B)


# ==========================================
# 🔹 DOCX Helpers
# ==========================================
def get_block_items(doc: Document) -> List[Tuple[str, object]]:
    blocks = []
    para_map = {p._p: p for p in doc.paragraphs}
    tbl_map = {t._tbl: t for t in doc.tables}
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            para = para_map.get(child)
            if para is not None:
                blocks.append(("p", para))
        elif isinstance(child, CT_Tbl):
            tbl = tbl_map.get(child)
            if tbl is not None:
                blocks.append(("tbl", tbl))
    return blocks

def detect_header_row(table, target_headers: List[str]) -> int:
    target_tokens = [tokens(h) for h in target_headers if h]
    best_idx = 0
    best_score = 0.0
    for r_idx in range(min(5, len(table.rows))):
        row_texts = [normalize(cell.text) for cell in table.rows[r_idx].cells]
        row_tokens = [tokens(t) for t in row_texts]
        if not target_tokens:
            continue
        scores = []
        for tt in target_tokens:
            best = 0.0
            for rt in row_tokens:
                s = jaccard(tt, rt)
                best = max(best, s)
            scores.append(best)
        avg = sum(scores) / len(scores) if scores else 0.0
        if avg > best_score:
            best_score = avg
            best_idx = r_idx
    return best_idx if best_score >= 0.15 else 0

def clear_table_body_after(table, header_row_idx: int):
    while len(table.rows) > header_row_idx + 1:
        tr = table.rows[header_row_idx + 1]._tr
        table._tbl.remove(tr)

def add_aligned_row(table, header_count: int, values: List[str]):
    new_cells = table.add_row().cells
    for i in range(header_count):
        v = values[i] if i < len(values) else ""
        new_cells[i].text = str(v)
        for p in new_cells[i].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)


def extract_table_with_multirow(tbl: Dict[str, Any], headers: List[str], vectorstore) -> ExtractedTable:
    """
    Simplified hybrid RAG-based table extraction.
    Requests the entire table in one structured query using LLM context grounding.
    """

    print("\n📊 === TABLE EXTRACTION START (Simplified) ===")
    table_name = tbl.get("name", "Unknown Table")
    print(f"🪶 Table Name: {table_name}")
    print(f"🎯 Headers: {headers}")

    # -----------------------------
    # STEP 1: Ask the LLM for the entire table directly
    # -----------------------------
    query = (
        f"Provide all rows for table '{table_name}' with headers: {headers}. "
        f"Use prior document context. Ensure numeric, date, and textual values "
        f"match the source document exactly. Include multiple rows if present."
    )

    print("\n🔍 Querying LLM for structured table extraction...",query)

    try:
        structured = rag(query, vectorstore, structure=ExtractedTable)
        print(f"🧠 RAG Structured Table Response: {structured}")

        # ✅ normalize whether it's dict or model
        # if isinstance(structured, dict):
        #     structured = ExtractedTable.model_validate(structured)

        # if structured and getattr(structured, "rows", None):
        #     print(f"✅ Structured extraction succeeded with {len(structured.rows)} rows.")
        return structured
    except Exception as e:
        print("⚠️ Structured extraction failed:", e)
        print(traceback.format_exc())


 
    # -----------------------------
    # STEP 2: Fallback if LLM didn’t return structured table
    # -----------------------------
    print("⚙️ Using fallback empty table.")
    table_data = ExtractedTable(
        name=table_name,
        headers=headers or [],
        rows=[],
        confidence=0.3,
        source=f"chunk_{random.randint(1,5)}.txt"
    )

    print(f"📦 Returning fallback ExtractedTable(name={table_data.name}, headers={table_data.headers}, rows={len(table_data.rows)})")
    print("📊 === TABLE EXTRACTION END ===\n")

    return table_data


# ==========================================
# 🔹 DOC FILLER MAIN FUNCTION (with reason + JSON)
# ==========================================
# def fill_docx_using_enriched(template_path: str, enriched_json_path: str, output_path: str = None) -> Dict[str, Any]:
async def fill_docx_using_enriched(template_path: str, enriched_json: any, output_path: str = None) -> Dict[str, Any]:
    # with open(enriched_json_path, "r", encoding="utf-8") as f:
    #     enriched = json.load(f)
    enriched = enriched_json

    retriever = create_retriever("./chroma_db_compliance3", "surveys_lease_data")
    vectorstore = retriever.vectorstore

    template = Path(template_path)
    if output_path is None:
        output_path = str(template.with_name(template.stem + "_filled" + template.suffix))
    json_output_path = str(template.with_name(template.stem + "_filled_data.json"))

    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    gathered_info = {"fields": [], "tables": []}

    # ===== TEXT FIELDS =====
    replacements: Dict[str, str] = {}
    for page in enriched.get("pages", []):
        for field in page.get("fields", []):
            name = field.get("name") or field.get("field_text") or ""
            key = name.rstrip(":").strip()

            if not field.get("value"):
                try:
                    # res = rag("value for "+field.get("query") or key, vectorstore, structure=ExtractedField)
                    res = await query_rag(field.get("query") or key, structured_schema=ExtractedField)
                except:
                    # res = rag("value for "+ field.get("name") or key, vectorstore, structure=ExtractedField)
                    res = await query_rag(field.get("query") or key, structured_schema=ExtractedField)

                # Normalize output (handle dict or model)
                if isinstance(res, dict):
                    field["value"] = res.get("value", "")
                    field["confidence"] = res.get("confidence", 0)
                    field["source"] = res.get("source", "")
                    field["reason"] = res.get("reason", "Not provided")
                else:
                    field["value"] = getattr(res, "value", "")
                    field["confidence"] = getattr(res, "confidence", 0)
                    field["source"] = getattr(res, "source", "")
                    field["reason"] = getattr(res, "reason", "Not provided")


            replacements[key] = f"{name} {field.get('value', '')}"
            replacements[f"{{{{{key}}}}}"] = field.get("value", "")
            replacements[f"{{{key}}}"] = field.get("value", "")

            gathered_info["fields"].append({
                "name": key,
                "value": field["value"],
                "confidence": field.get("confidence", 0),
                "source": field.get("source", ""),
                "reason": field.get("reason", "")
            })

    for para in doc.paragraphs:
        text = para.text
        for old, new in replacements.items():
            if old in text:
                text = text.replace(old, new)
        para.text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    for old, new in replacements.items():
                        if old in text:
                            text = text.replace(old, new)
                    para.text = text

    # ===== TABLES =====
    blocks = get_block_items(doc)
    for page in enriched.get("pages", []):
        print("page iteration ")
        for tbl in page.get("tables", []):
            print("table iteration ")
            target_headers = tbl.get("headers") or [] or (list(tbl.get("rows", [])[0].keys()) if tbl.get("rows") else [])
            if not target_headers:
                continue

            label = normalize(tbl.get("name", ""))
            label_block_index = None
            for idx, (k, obj) in enumerate(blocks):
                if k == 'p' and label and label in normalize(obj.text):
                    label_block_index = idx
                    break

            candidate_table = None
            best_dist = None
            for idx, (k, obj) in enumerate(blocks):
                if k == 'tbl':
                    if label_block_index is not None:
                        dist = abs(idx - label_block_index)
                        if best_dist is None or dist < best_dist:
                            best_dist = dist
                            candidate_table = obj
                    else:
                        candidate_table = obj

            if not candidate_table:
                print(f"⚠️ No matching table found for '{tbl.get('name')}'")
                continue

            header_row_idx = detect_header_row(candidate_table, target_headers)
            clear_table_body_after(candidate_table, header_row_idx)

            extracted = extract_table_with_multirow(tbl, target_headers, vectorstore)
            # normalize structured or dict output
            if isinstance(extracted, dict):
                extracted_rows = extracted.get("rows", [])
                confidence = extracted.get("confidence", 0.0)
                source = extracted.get("source", "")
            else:
                extracted_rows = getattr(extracted, "rows", [])
                confidence = getattr(extracted, "confidence", 0.0)
                source = getattr(extracted, "source", "")

            # ensure source is string
            if isinstance(source, list):
                source = ", ".join(source)

            # build table rows
            for row in extracted_rows:
                header_cell_count = len(candidate_table.rows[header_row_idx].cells)

                # if row is list, use directly; if dict, map by headers
                if isinstance(row, list):
                    vals = row
                elif isinstance(row, dict):
                    vals = [row.get(h, "") for h in target_headers]
                else:
                    vals = []

                # pad or trim values to match table width
                vals = vals[:header_cell_count] + [""] * (header_cell_count - len(vals))
                add_aligned_row(candidate_table, header_cell_count, vals)

            # append normalized table info
            gathered_info["tables"].append({
                "name": tbl.get("name"),
                "headers": target_headers,
                "rows": extracted_rows,
                "confidence": confidence,
                "source": source
            })

    doc.save(output_path)

    with open(json_output_path, "w", encoding="utf-8") as jf:
        json.dump(gathered_info, jf, indent=2)

    print(f"📘 Filled document written to: {output_path}")
    print(f"📦 JSON data saved to: {json_output_path}")

    return {"docx_path": output_path, "json_path": json_output_path, "data": gathered_info}


if __name__ == "__main__":
    TEMPLATE = "userPdfData/refinalfollowupactionrequiredforyourprototypedeve (1)/Lease Abstract Template.docx"
    ENRICHED = "enriched_structure.json"
    out = fill_docx_using_enriched(TEMPLATE, ENRICHED)
    print("✅ Done ->", out)
