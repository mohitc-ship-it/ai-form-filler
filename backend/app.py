from fastapi import FastAPI, UploadFile, File, Request, Query, Form
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional
from pathlib import Path
import tempfile
import shutil
import json
from fill import fill_excel , fill_docx, reqHandler
import uuid 
# Libraries for document generation
from openpyxl import Workbook
from docx import Document
from datetime import datetime

app = FastAPI(title="AI Form Filler API")

# --- Enable CORS ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Static files folder ---
STATIC_DIR = Path(__file__).parent / "static"
STATIC_DIR.mkdir(exist_ok=True)

# --- Fallback demo data ---
DEFAULT_DOCX_DATA = [
    {"name": "Tenant Name", "value": "Mid State Plaza"},
    {"name": "Square Feet", "value": "12000"},
    {"name": "Term", "value": "5 years"},
    {"name": "Rent", "value": "$5000/month"},
]

DEFAULT_XLSX_DATA = [
    {"name": "Company", "value": "ConsultAdd Inc."},
    {"name": "Address", "value": "123 AI Street"},
    {"name": "City", "value": "New York"},
    {"name": "Zip Code", "value": "10001"},
]

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_ROOT = BASE_DIR / "uploaded_files"
UPLOAD_ROOT.mkdir(exist_ok=True)  # ensure folder exists

# Optional: point to your static/fallback JSON data
STATIC_DIR = BASE_DIR / "static"

@app.post("/api/process")
async def process_files_dynamic(
    files: Optional[List[UploadFile]] = File(None),
    files_meta: Optional[str] = Form(None),
):
    """
    Dynamic version of /api/process:
    - Saves uploaded files into uploaded_files/{session_id}/context & forms
    - Calls reqHandler() to process dynamically (fill_docx / fill_excel)
    - Returns JSON with session info and dynamic results
    """
    try:
        # --- 1️⃣ Validate files ---
        if not files:
            return JSONResponse({"error": "No files uploaded"}, status_code=400)

        # --- 2️⃣ Create session directories ---
        session_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        session_dir = UPLOAD_ROOT / f"{timestamp}_{session_id}"
        context_dir = session_dir / "context"
        forms_dir = session_dir / "forms"

        context_dir.mkdir(parents=True, exist_ok=True)
        forms_dir.mkdir(parents=True, exist_ok=True)

        print(f"✅ Dynamic processing started for session: {session_id}")
        print(f"📂 Files will be saved in: {session_dir}")

        # --- 3️⃣ Handle metadata ---
        meta_list = []
        if files_meta:
            try:
                meta_list = json.loads(files_meta)
            except Exception as e:
                print("❌ Failed to parse files_meta JSON:", e)

        saved_files = []
        files_lookup = {f.filename: f for f in (files or [])}

        # --- 4️⃣ Save uploaded files into correct folders ---
        for m in meta_list:
            name = m.get("file_name") or m.get("file_path") or m.get("filename")
            if not name:
                continue

            upload_file = files_lookup.get(name)
            if not upload_file:
                continue

            ftype = m.get("file_type", "").lower()
            target_dir = forms_dir if ftype == "form" else context_dir
            file_path = target_dir / upload_file.filename

            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(upload_file.file, buffer)

            saved_files.append({"meta": m, "path": str(file_path)})

        # fallback: save all if no metadata
        if not meta_list:
            for f in files:
                ftype = "form" if f.filename.lower().endswith((".docx", ".xlsx")) else "context"
                target_dir = forms_dir if ftype == "form" else context_dir
                file_path = target_dir / f.filename
                with open(file_path, "wb") as buffer:
                    shutil.copyfileobj(f.file, buffer)
                saved_files.append({
                    "meta": {"file_type": ftype, "file_name": f.filename},
                    "path": str(file_path)
                })

        # --- 5️⃣ Run dynamic request handler (main logic) ---
        print(f"⚙️ Running reqHandler for session {session_id} ...")
        result = await reqHandler(session_id, saved_files)

        # --- 6️⃣ Return dynamic processing result ---
        return JSONResponse({
            "session_id": session_id,
            "saved_dir": str(session_dir),
            "files": saved_files,
            "result": str(result)
        })

    except Exception as e:
        print("❌ Error in /api/process:", e)
        return JSONResponse({"error": str(e)}, status_code=500)


# @app.post("/api/process")
# async def process_files(
#     files: Optional[List[UploadFile]] = File(None),
#     files_meta: Optional[str] = Form(None),
# ):
#     """
#     Detect form type (.docx or .xlsx),
#     store uploaded files inside project repo (with 'context' and 'forms' subfolders),
#     and return structured JSON output (static for now).
#     """

#     try:
#         # --- 1️⃣ Basic checks ---
#         if not files:
#             return JSONResponse({"error": "No files uploaded"}, status_code=400)

#         # --- 2️⃣ Create unique session folder ---
#         session_id = str(uuid.uuid4())
#         timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#         session_dir = UPLOAD_ROOT / f"{timestamp}_{session_id}"
#         context_dir = session_dir / "context"
#         forms_dir = session_dir / "forms"

#         # Create folders
#         context_dir.mkdir(parents=True, exist_ok=True)
#         forms_dir.mkdir(parents=True, exist_ok=True)

#         print(f"✅ Files saved to: {session_dir}")

#         # --- 3️⃣ Parse metadata and save uploaded files ---
#         meta_list = []
#         if files_meta:
#             try:
#                 meta_list = json.loads(files_meta)
#             except Exception as e:
#                 print("❌ Failed to parse files_meta JSON:", e)

#         saved_files = []
#         files_lookup = {f.filename: f for f in (files or [])}

#         for m in meta_list:
#             name = m.get("file_name") or m.get("file_path") or m.get("filename")
#             if not name:
#                 continue

#             upload_file = files_lookup.get(name)
#             if not upload_file:
#                 print(f"⚠️ Metadata references unknown file: {name}")
#                 continue

#             ftype = m.get("file_type", "").lower()
#             target_dir = forms_dir if ftype == "form" else context_dir
#             file_path = target_dir / upload_file.filename

#             with open(file_path, "wb") as buffer:
#                 shutil.copyfileobj(upload_file.file, buffer)

#             saved_files.append({"meta": m, "path": str(file_path)})

#         # If no metadata, save all uploaded files heuristically
#         if not meta_list and files:
#             for f in files:
#                 ftype = "form" if f.filename.lower().endswith((".docx", ".xlsx")) else "context"
#                 target_dir = forms_dir if ftype == "form" else context_dir
#                 file_path = target_dir / f.filename

#                 with open(file_path, "wb") as buffer:
#                     shutil.copyfileobj(f.file, buffer)

#                 saved_files.append({
#                     "meta": {"file_type": ftype, "file_name": f.filename},
#                     "path": str(file_path)
#                 })

#         # --- 4️⃣ Determine which saved file is the form (from metadata) ---
#         form_filename = None
#         for s in saved_files:
#             meta = s.get("meta", {})
#             if meta.get("file_type") == "form":
#                 form_filename = meta.get("file_name")
#                 break

#         # fallback: first docx/xlsx
#         if not form_filename:
#             for s in saved_files:
#                 p = s.get("path", "")
#                 if p.lower().endswith((".docx", ".xlsx")):
#                     form_filename = Path(p).name
#                     break

#         filename = (form_filename or "").lower()

#         # --- 5️⃣ Load corresponding static JSON (temporary mock logic) ---
#         if filename.endswith(".docx"):
#             json_path = STATIC_DIR / "docsxx.json"
#         elif filename.endswith(".xlsx"):
#             json_path = STATIC_DIR / "excels.json"
#         else:
#             json_path = None

#         print("json path:", json_path)

#         if json_path and json_path.exists():
#             with open(json_path, "r") as f:
#                 data = json.load(f)
#         else:
#             data = DEFAULT_DOCX_DATA if filename.endswith(".docx") else DEFAULT_XLSX_DATA
        



        

#         # ✅ Return mock response
#         return JSONResponse({
#             "session_id": session_id,
#             "saved_dir": str(session_dir),
#             "files": saved_files,
#             "data": data
#         })
        
#         reqHandler(session_id,files);

#         # return JSONResponse(data);
#         # ----------------------------------------------------------
#         # 🧩 Future Implementation (enable once ready):
#         # ----------------------------------------------------------
#         # if filename.endswith(".docx"):
#         #     result = await fill_docx(form_file, context_files)
#         # elif filename.endswith(".xlsx"):
#         #     result = await fill_excel(form_file, context_files)
#         # return JSONResponse({
#         #     "session_id": session_id,
#         #     "saved_dir": str(session_dir),
#         #     "results": result
#         # })
#         # ----------------------------------------------------------

#     except Exception as e:
#         print("❌ Error in /api/process:", e)
#         return JSONResponse({"error": str(e)}, status_code=500)

# 🧾 STEP 2: Download filled DOCX or XLSX
# @app.post("/api/download")
# async def download_filled_file(
#     request: Request,
#     type: str = Query(..., regex="^(docx|xlsx)$"),
#     session_id: Optional[str] = Query(None),
# ):
#     """
#     Convert JSON data from frontend into a filled .docx or .xlsx file
#     and return it as download.
#     """
#     try:
#         body = await request.json()
#         results = body.get("results", [])

#         tmp_dir = Path(tempfile.mkdtemp())

#         # If a session_id is provided, try to find any file under known
#         # result folders that contains the session_id in its filename or path.
#         def find_session_file(session: str, ext: str) -> Optional[Path]:
#             # candidate directories to search for generated results
#             candidates = [UPLOAD_ROOT, BASE_DIR / "marker-output_folder", BASE_DIR / "static"]
#             for base in candidates:
#                 if not base.exists():
#                     continue
#                 for p in base.rglob(f"*{session}*{ext}"):
#                     return p
#             return None

#         # Try to locate a session-specific file if session_id provided
#         session_file_path: Optional[Path] = None
#         if session_id:
#             session_file_path = find_session_file(session_id, f".{type}")

#         if type == "docx":
#             # doc = Document()
#             # doc.add_heading("Filled DOCX Form", level=1)
#             # for r in results:
#             #     doc.add_paragraph(f"{r.get('name', '')}: {r.get('value', '')}")
#             # file_path = tmp_dir / "filled_form.docx"
#             # doc.save(file_path)
#             if session_file_path and session_file_path.exists():
#                 return FileResponse(
#                     str(session_file_path),
#                     media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                     filename=session_file_path.name,
#                 )

#             file_path = BASE_DIR / "static" / "Lease Abstract Template_filled.docx"
#             return FileResponse(
#                 str(file_path),
#                 media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                 filename="Lease Abstract Template_filled.docx",
#             )

#         elif type == "xlsx":
#             # wb = Workbook()
#             # ws = wb.active
#             # ws.title = "Filled Form"
#             # ws.append(["Field", "Value"])
#             # for r in results:
#             #     ws.append([r.get("name", ""), r.get("value", "")])
#             # file_path = tmp_dir / "filled_form.xlsx"
#             # wb.save(file_path)
#             if session_file_path and session_file_path.exists():
#                 return FileResponse(
#                     str(session_file_path),
#                     media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
#                     filename=session_file_path.name,
#                 )

#             file_path = BASE_DIR / "static" / "BOV Template_filled.xlsx"
#             return FileResponse(
#                 str(file_path),
#                 media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
#                 filename="BOV Template_filled.xlsx",
#             )

#     except Exception as e:
#         print("❌ Error in /api/download:", e)
#         return JSONResponse({"error": str(e)}, status_code=500)



# # IMPLEMENTATION FROM ACTUAL UPLOAD FILE SAVING AND CHECKING IF ALREADY EXISTS



@app.post("/api/download")
async def download_filled_file(
    request: Request,
    type: str = Query(..., regex="^(docx|xlsx)$"),
    session_id: Optional[str] = Query(None),
):
    """
    Locate and return a filled .docx or .xlsx file for a given session.
    Searches in uploaded_files/{session_id}/forms/ before falling back to static defaults.
    """
    try:
        body = await request.json()
        results = body.get("results", [])

        tmp_dir = Path(tempfile.mkdtemp())

        def find_session_file(session: str, ext: str) -> Optional[Path]:
            """
            Searches inside uploaded_files/{session_id}/forms/ for any file
            containing both session_id and 'filled' in its name.
            """
            base_folder = Path(UPLOAD_ROOT) / session / "forms"
            if not base_folder.exists():
                print(f"⚠️ Folder not found: {base_folder}")
                return None

            print(f"🔍 Searching for filled form in: {base_folder}")
            for p in base_folder.glob(f"*{session}*filled*{ext}"):
                print(f"✅ Found file: {p}")
                return p

            print(f"⚠️ No filled {ext} file found for session {session}")
            return None

        session_file_path: Optional[Path] = None
        if session_id:
            session_file_path = find_session_file(session_id, f".{type}")

        # ------------------------------------------
        # DOCX HANDLING
        # ------------------------------------------
        if type == "docx":
            if session_file_path and session_file_path.exists():
                return FileResponse(
                    str(session_file_path),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    filename=session_file_path.name,
                )

            # fallback if not found
            default_path = BASE_DIR / "static" / "Lease Abstract Template_filled.docx"
            print(f"⚙️ Using fallback DOCX: {default_path}")
            return FileResponse(
                str(default_path),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename="Lease Abstract Template_filled.docx",
            )

        # ------------------------------------------
        # XLSX HANDLING
        # ------------------------------------------
        elif type == "xlsx":
            if session_file_path and session_file_path.exists():
                return FileResponse(
                    str(session_file_path),
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    filename=session_file_path.name,
                )

            # fallback if not found
            default_path = BASE_DIR / "static" / "BOV Template_filled.xlsx"
            print(f"⚙️ Using fallback XLSX: {default_path}")
            return FileResponse(
                str(default_path),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                filename="BOV Template_filled.xlsx",
            )

    except Exception as e:
        print("❌ Error in /api/download:", e)
        return JSONResponse({"error": str(e)}, status_code=500)
